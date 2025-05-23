import os
import docx
import traceback
from pdfminer.high_level import extract_text as pdfminer_extract_text

def process_document(file_path):
    """
    Extract text from PDF or Word document
    
    Args:
        file_path (str): Path to the document file
        
    Returns:
        str: Extracted text from the document
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        print(traceback.format_exc())
        raise Exception(f"Failed to process document: {str(e)}")

def extract_text_from_pdf(pdf_path):
    """
    Extract text from PDF file using pdfminer.six
    
    Args:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF
    """
    try:
        text = pdfminer_extract_text(pdf_path)
        if not text.strip():
            raise ValueError("PDF does not contain extractable text. It may be scanned or image-based.")
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        print(traceback.format_exc())
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_path):
    """
    Extract text from Word document
    
    Args:
        docx_path (str): Path to the Word document
        
    Returns:
        str: Extracted text from the Word document
    """
    try:
        doc = docx.Document(docx_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
        
        # If document has no text, raise error
        if not text.strip():
            raise ValueError("Word document does not contain any text")
            
        return text
    except Exception as e:
        print(f"Error extracting text from Word document: {str(e)}")
        print(traceback.format_exc())
        raise ValueError(f"Failed to process Word document: {str(e)}") 